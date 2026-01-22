"""
Word Document Generator Module
Generates editable Word documents from extracted text.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional, Tuple
import re
import io


class WordGenerator:
    """Generates Word documents from extracted OCR text."""
    
    def __init__(self):
        """Initialize the Word generator."""
        self.doc = None
    
    def create_document(self) -> Document:
        """Create a new Word document."""
        self.doc = Document()
        
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        return self.doc
    
    def _is_header(self, text: str) -> bool:
        """
        Determine if text appears to be a header/title.
        
        Args:
            text: Text to check
            
        Returns:
            True if text appears to be a header
        """
        # Headers are often:
        # - Short (less than 50 chars)
        # - Start with # or numbers
        # - All caps or title case
        # - Contains keywords like "Chapter", "Section", etc.
        
        text = text.strip()
        if not text:
            return False
        
        # Check for common header patterns
        header_patterns = [
            r'^#\s*',  # Starts with #
            r'^\d+\.\s*[A-Z]',  # Numbered section
            r'^[A-Z][A-Z\s]+$',  # All caps
            r'^(Chapter|Section|Part|Unit)\s+\d*',  # Common header words
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        # Short text that ends with colon
        if len(text) < 50 and text.endswith(':'):
            return True
        
        return False
    
    def _detect_formatting(self, text: str) -> Dict:
        """
        Detect special formatting needs for text.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with formatting hints
        """
        formatting = {
            'is_header': self._is_header(text),
            'is_bullet': text.strip().startswith(('•', '-', '*', '○', '·')),
            'is_numbered': bool(re.match(r'^\d+[\.\)]\s', text.strip())),
            'has_equation': bool(re.search(r'[=+\-×÷∑∫√∞≠≤≥]', text)),
        }
        return formatting
    
    def add_paragraph(self, text: str, is_header: bool = False, 
                      color: Optional[Tuple[int, int, int]] = None) -> None:
        """
        Add a paragraph to the document.
        
        Args:
            text: Text content
            is_header: Whether this is a header
            color: Optional RGB color tuple
        """
        if not self.doc:
            self.create_document()
        
        if is_header:
            para = self.doc.add_heading(text, level=2)
        else:
            para = self.doc.add_paragraph(text)
        
        if color and para.runs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(*color)
    
    def add_two_column_content(self, left_texts: List[Dict], 
                                 right_texts: List[Dict]) -> None:
        """
        Add content in two-column layout using a table.
        
        Args:
            left_texts: List of text dictionaries for left column
            right_texts: List of text dictionaries for right column
        """
        if not self.doc:
            self.create_document()
        
        # Create a table with 1 row and 2 columns
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(3.25)
        table.columns[1].width = Inches(3.25)
        
        # Get cells
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # Add content to left column
        self._add_texts_to_cell(left_cell, left_texts)
        
        # Add content to right column
        self._add_texts_to_cell(right_cell, right_texts)
        
        # Remove table borders for cleaner look
        self._remove_table_borders(table)
    
    def _add_texts_to_cell(self, cell, texts: List[Dict]) -> None:
        """
        Add text content to a table cell.
        
        Args:
            cell: Table cell object
            texts: List of text dictionaries
        """
        # Clear default paragraph
        cell.paragraphs[0].clear()
        
        # Group texts into lines
        from .ocr_processor import OCRProcessor
        processor = OCRProcessor.__new__(OCRProcessor)
        lines = processor.group_into_lines(texts)
        
        first = True
        for line in lines:
            line_text = ' '.join([t['text'] for t in line])
            formatting = self._detect_formatting(line_text)
            
            if first:
                para = cell.paragraphs[0]
                first = False
            else:
                para = cell.add_paragraph()
            
            run = para.add_run(line_text)
            
            if formatting['is_header']:
                run.bold = True
                run.font.size = Pt(12)
                # Red color for headers (matching the original handwritten notes)
                run.font.color.rgb = RGBColor(180, 0, 0)
    
    def _remove_table_borders(self, table) -> None:
        """Remove all borders from a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def generate_from_column_texts(self, column_texts: Dict[int, List[Dict]], 
                                    title: Optional[str] = None) -> Document:
        """
        Generate a Word document from column-organized text.
        
        Args:
            column_texts: Dictionary mapping column index to text dictionaries
            title: Optional document title
            
        Returns:
            Generated Document object
        """
        self.create_document()
        
        if title:
            self.doc.add_heading(title, level=1)
        
        num_columns = len(column_texts)
        
        if num_columns == 2:
            # Two-column layout
            left_texts = column_texts.get(0, [])
            right_texts = column_texts.get(1, [])
            self.add_two_column_content(left_texts, right_texts)
        else:
            # Single column or more - add sequentially
            for col_idx in sorted(column_texts.keys()):
                texts = column_texts[col_idx]
                
                # Group into lines
                from .ocr_processor import OCRProcessor
                processor = OCRProcessor.__new__(OCRProcessor)
                lines = processor.group_into_lines(texts)
                
                for line in lines:
                    line_text = ' '.join([t['text'] for t in line])
                    formatting = self._detect_formatting(line_text)
                    
                    if formatting['is_header']:
                        para = self.doc.add_heading(line_text, level=2)
                        if para.runs:
                            para.runs[0].font.color.rgb = RGBColor(180, 0, 0)
                    else:
                        self.doc.add_paragraph(line_text)
                
                # Add separator between columns
                if col_idx < max(column_texts.keys()):
                    self.doc.add_paragraph('─' * 50)
        
        return self.doc
    
    def generate_from_text(self, text: str, title: Optional[str] = None) -> Document:
        """
        Generate a Word document from plain text.
        
        Args:
            text: Plain text content
            title: Optional document title
            
        Returns:
            Generated Document object
        """
        self.create_document()
        
        if title:
            self.doc.add_heading(title, level=1)
        
        # Split into paragraphs
        paragraphs = text.split('\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            formatting = self._detect_formatting(para_text)
            
            if formatting['is_header']:
                para = self.doc.add_heading(para_text, level=2)
                if para.runs:
                    para.runs[0].font.color.rgb = RGBColor(180, 0, 0)
            elif para_text == '---':
                # Column separator
                self.doc.add_paragraph('─' * 50)
            else:
                self.doc.add_paragraph(para_text)
        
        return self.doc
    
    def save(self, filepath: str) -> str:
        """
        Save the document to a file.
        
        Args:
            filepath: Path to save the document
            
        Returns:
            Path to saved file
        """
        if not self.doc:
            raise ValueError("No document created. Call create_document() first.")
        
        self.doc.save(filepath)
        return filepath
    
    def save_to_bytes(self) -> bytes:
        """
        Save the document to bytes (for web download).
        
        Returns:
            Document as bytes
        """
        if not self.doc:
            raise ValueError("No document created. Call create_document() first.")
        
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
